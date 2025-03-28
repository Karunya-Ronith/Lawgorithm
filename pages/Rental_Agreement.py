import streamlit as st
import requests
from docx import Document
import io

# Page Title
st.title("🏠 Indian Rental Agreement Generator")

st.markdown("""
    <style>
        .stTextInput, .stNumberInput, .stTextArea, .stDateInput {
            background-color: #1E1E1E !important;
            color: #E0E0E0 !important;
            border-radius: 8px !important;
        }
        .stForm {
            background-color: #2C2C2C; 
            padding: 20px; 
            border-radius: 10px; 
            box-shadow: 2px 2px 10px rgba(255, 255, 255, 0.1);
        }
        .stButton>button {
            background-color: #BB86FC !important;
            color: white !important;
            font-size: 16px;
            border-radius: 8px;
            padding: 10px;
        }
    </style>
""", unsafe_allow_html=True)

# ---- Rental Agreement Form ----
with st.form("rental_agreement_form"):
    st.subheader("📝 Rental Agreement Details")

    # Tenant & Owner Information
    with st.expander("👤 Parties Involved"):
        tenant_name = st.text_input("Tenant Name")
        tenant_state = st.text_input("Tenant State")
        owner_name = st.text_input("Owner Name")
        owner_state = st.text_input("Owner State")

    # Witnesses & Date
    with st.expander("🗓️ Witnesses & Signing Details"):
        witness1_name = st.text_input("Witness 1 Name")
        witness2_name = st.text_input("Witness 2 Name")
        signing_date = st.date_input("Signing Date")
        signing_day = st.text_input("Signing Day (Monday, Tuesday, etc.)")
        signing_time = st.text_input("Signing Time (HH:MM AM/PM)")
        signing_month = st.text_input("Signing Month")

    # Property & Rental Details
    with st.expander("🏠 Property & Rental Terms"):
        land_location = st.text_area("Location of the Property")
        rental_contract_years = st.number_input("Rental Contract Duration (Years)", min_value=1, step=1)
        start_year = st.number_input("Start Year of Agreement", min_value=1900, step=1)
        start_month = st.text_input("Rental Agreement Starting Month")
        monthly_rent = st.number_input("Monthly Rent (₹)", min_value=500, step=100)
        payment_deadline = st.number_input("Payment Deadline (Months)", min_value=1, step=1)

    # Financial Terms
    with st.expander("💰 Financial Terms"):
        interest_rate = st.number_input("Rate of Interest (%)", min_value=0.0, step=0.1)
        land_tax = st.number_input("Tax on Property (₹)", min_value=0.0, step=100.0)

    submit_button = st.form_submit_button("⚖️ Generate Legal Agreement")

# ---- Processing & Generating Agreement ----
if submit_button:
    prompt = f"""
    You are an Indian legal expert. Generate a legally valid **rental agreement** based on the following details:

    **Parties Involved:**
    - Tenant: {tenant_name} ({tenant_state})
    - Owner: {owner_name} ({owner_state})
    - Witnesses: {witness1_name}, {witness2_name}

    **Signing Details:**
    - Date: {signing_date} ({signing_day} at {signing_time}, {signing_month})

    **Property & Rental Terms:**
    - Location: {land_location}
    - Duration: {rental_contract_years} years (Starting {start_month}, {start_year})
    - Monthly Rent: ₹{monthly_rent}
    - Payment Deadline: {payment_deadline} months

    **Financial Terms:**
    - Interest Rate: {interest_rate}%
    - Property Tax: ₹{land_tax}

    Ensure the document follows **Indian rental laws** and includes clauses for **rent payment, termination, obligations, legal jurisdiction, and penalties**. Write in **formal legal language**.
    """

    # Call Mistral API
    ollama_response = requests.post(
        "http://localhost:11434/api/generate",
        json={"model": "mistral", "prompt": prompt, "stream": False}
    ).json()

    rental_agreement = ollama_response.get("response", "Sorry, I couldn't generate the legal agreement.")

    # ---- Display Generated Agreement ----
    st.subheader("📜 Generated Rental Agreement")
    st.text_area("", rental_agreement, height=300)

    # ---- Save & Download as Word Document ----
    doc = Document()
    doc.add_heading("Rental Agreement", level=1)
    doc.add_paragraph(rental_agreement)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    st.download_button(label="📄 Download Rental Agreement",
                       data=file_stream,
                       file_name="Rental_Agreement.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.success("✅ Rental Agreement Generated Successfully!")
