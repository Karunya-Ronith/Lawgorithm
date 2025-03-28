import streamlit as st
import pytesseract
from PIL import Image
from googletrans import Translator
from docx import Document
import io

# Set Tesseract OCR path (modify if necessary)
pytesseract.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract.exe'

# Initialize Google Translator
translator = Translator()

def extract_and_translate(image):
    try:
        # Extract Hindi text using Tesseract OCR
        hindi_text = pytesseract.image_to_string(image, lang="hin")
        
        # Translate Hindi text to English
        translated_text = translator.translate(hindi_text, src="hi", dest="en").text
        
        return hindi_text, translated_text
    except Exception as e:
        st.error(f"Error processing image: {e}")
        return None, None

st.title("📝 Hindi to English Document Translator")

uploaded_file = st.file_uploader("Upload an image containing Hindi text", type=["png", "jpg", "jpeg"])
 
if uploaded_file is not None:
    image = Image.open(uploaded_file)
    st.image(image, caption="Uploaded Image", use_container_width=True)
    
    with st.spinner("Extracting and translating..."):
        hindi_text, translated_text = extract_and_translate(image)
        
    if translated_text:
        st.subheader("Extracted Hindi Text:")
        st.text_area("", hindi_text, height=150)
        
        st.subheader("Translated English Text:")
        st.text_area("", translated_text, height=150)
        
        # Create Word document
        doc = Document()
        doc.add_paragraph("**Extracted Hindi Text:**")
        doc.add_paragraph(hindi_text)
        doc.add_paragraph("\n**Translated English Text:**")
        doc.add_paragraph(translated_text)
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        st.download_button(
            "📄 Download Translated Document", file_stream, "Translated_Text.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
